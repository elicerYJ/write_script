import streamlit as st
import pandas as pd
import openpyxl  # Excel 파일로 저장하기 위해 필요
from io import BytesIO  # 파일 다운로드를 위해 필요
from docx import Document  # 워드 파일로 저장하기 위해 필요

def clean_text(text):
    """XML 호환 문자만 남기고 모두 제거 또는 대체"""
    return ''.join(char for char in text if char.isprintable() or char in '\n\t').strip()

# 데이터프레임을 Excel 파일로 변환하는 함수
def to_excel(df):
    output = BytesIO()
    
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        df.to_excel(writer, index=False, sheet_name='Sheet1')
    
    # 파일 데이터를 가져와서 반환
    return output.getvalue()



def to_word(df):
    doc = Document()

    # 데이터프레임의 각 행을 반복 처리
    for index, row in df.iterrows():
        doc.add_paragraph(f"페이지 번호 : {clean_text(str(row['페이지 번호']))}")
        doc.add_paragraph(f"애니메이션 적용 대상 : {clean_text(str(row['애니메이션 적용 대상']))}")
        doc.add_paragraph(f"애니메이션 효과 설명 : {clean_text(str(row['애니메이션 효과 설명']))}")
        doc.add_paragraph("<대본>")
        doc.add_paragraph(clean_text(str(row['사용할 대본'])))
        doc.add_paragraph("---")  # 구분자 추가

    # 파일 데이터를 BytesIO 객체로 저장하고 반환
    doc_io = BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    return doc_io.getvalue()


st.set_page_config(layout="wide")

main_col1, main_col2 = st.columns(2)


# 데이터프레임을 저장할 빈 데이터프레임 생성
if 'df' not in st.session_state:
    st.session_state.df = pd.DataFrame(columns=['페이지 번호', 
                                                '애니메이션 적용 대상', 
                                                '애니메이션 효과 설명',
                                                '사용할 대본'])

with main_col1 :
    # 사용자 입력 받기
    with st.form(key='record_form'):
        st.write("### 스크립트 작성")
        
        form_col1, form_col2 = st.columns([0.4,0.6])

        with form_col1 :
            course_name = st.text_input("강의명", placeholder="도레미 파이썬 vol1")
            page_number = st.number_input("페이지 번호", value=0, step=1)
                    
        with form_col2 :
            lecture_name = st.text_input("렉쳐명", placeholder="컴퓨터 과학이란?")
            animation_target = st.selectbox(
                '애니메이션 적용 대상',
                ('⛔ 없음', "🔠 텍스트", '🆚 도형을 포함한 텍스트', '🟪 도형', '🖼️ 이미지(아이콘)/코드', '✨ 효과', '📊 그래프', '👩‍🎨 애니메이션 제작 필요','🎸 기타'),
                index = 0
            )

        effect_comm = st.text_input('애니메이션 효과 설명', '없음')
        
        script = st.text_area('사용할 대본', height = 150)

        submit_button = st.form_submit_button(label='스크립트 추가')
        
        # 정보 추가 버튼이 클릭되었을 때
        if submit_button:
            # 입력받은 정보를 데이터프레임에 추가
            new_data = {'페이지 번호': page_number, 
                        '애니메이션 적용 대상': animation_target, 
                        '애니메이션 효과 설명' : effect_comm,
                        '사용할 대본': script}
            
            st.session_state.df = pd.concat(
                [st.session_state.df, pd.DataFrame([new_data])], 
                ignore_index=True
            )
    
    button_col1, button_col2, button_col3 = st.columns(3, gap='small')
    
    with button_col1 :
        if st.button('📜 워드 다운로드'):
            word_val = to_word(st.session_state.df)
            st.download_button(
                label='현재 데이터 워드로 다운로드', 
                data=word_val, 
                file_name=f"{course_name}_{lecture_name}_script.docx", 
                mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                )
    
    with button_col2 :
        if st.button('📋 엑셀 다운로드'):
            val = to_excel(st.session_state.df)
            st.download_button(
                label='현재 데이터 엑셀로 다운로드', 
                data=val, 
                file_name=f"{course_name}_{lecture_name}_script.xlsx", 
                mime='application/vnd.ms-excel'
                )
    
    with button_col3 :
        # 데이터프레임에서 마지막 행을 삭제하는 버튼
        if st.button('마지막 스크립트 삭제'):
            if not st.session_state.df.empty:
                st.session_state.df = st.session_state.df[:-1]
            else:
                st.warning('데이터프레임이 비어 있습니다.')


with main_col2 :
    st.write("### 슬라이트 노트")
    st.write("작성된 스크립트를 복사해서 슬라이드 노트에 붙여넣기 할 수 있습니다")
    copy_text = "애니메이션 적용 대상 : " + animation_target + " (" + effect_comm + ') \n' + "\n[대본]\n" + script
    
    st.code(
        copy_text,
        language='plain'
    )
    
    st.divider()
    
    # 데이터프레임을 화면에 표시
    #st.checkbox("Use container width", value=True, key="use_container_width")

    st.data_editor(
        st.session_state.df, 
        #use_container_width=st.session_state.use_container_width,
        width = 1000,
        hide_index=True
        )



st.divider()

# 데이터프레임 초기화 버튼
if st.button('스크립트 초기화'):
    st.session_state.df = pd.DataFrame(columns=['페이지 번호', 
                                                '애니메이션 적용 대상', 
                                                '애니메이션 효과 설명',
                                                '사용할 대본'])